""" docx

Word document parser.
"""

from typing import Optional
from docx import Document
from io import BytesIO

from app.modules.dataset.infra.parsers.base import DocumentParser, ParsedDocument


class DocxParser(DocumentParser):
    """Parser for Word (.docx) files."""
    
    async def parse(
        self,
        content: bytes,
        mime_type: Optional[str] = None,
        filename: Optional[str] = None,
    ) -> ParsedDocument:
        """Parse Word document content.
        
        Args:
            content: Word document content bytes.
            mime_type: Optional MIME type.
            filename: Optional filename.
            
        Returns:
            ParsedDocument instance.
        """
        doc_file = BytesIO(content)
        doc = Document(doc_file)
        
        # Extract text from all paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(paragraphs)
        
        # Extract metadata
        metadata = {
            "paragraph_count": len(paragraphs),
            "char_count": len(full_text),
        }
        
        # Try to extract title from first paragraph or core properties
        if paragraphs:
            metadata["title"] = paragraphs[0][:100]  # First paragraph as title
        
        if doc.core_properties.title:
            metadata["title"] = doc.core_properties.title
        
        return ParsedDocument(
            text=full_text,
            metadata=metadata,
        )
    
    def supports(self, mime_type: str) -> bool:
        """Check if supports MIME type."""
        return mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        )

